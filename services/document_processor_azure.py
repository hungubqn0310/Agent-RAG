import os
import io
import hashlib
from typing import List, Dict, Any, Tuple
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import openpyxl
import numpy as np
import json
from services.azure_openai_service import AzureOpenAIService
from langchain.text_splitter import RecursiveCharacterTextSplitter
import datetime
import shutil
from services.azure_document_intelligence import AzureDocumentIntelligenceService
import ocrmypdf
from pypdf import PdfReader

class DocumentProcessor:
    def __init__(self):
        self.azure_openai = AzureOpenAIService()
        self.batch_size = 1000  # Batch size cho embedding
        
        # Sử dụng RecursiveCharacterTextSplitter như mcp-rag
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len,
        )
        
        # Azure Document Intelligence (required for OCR)
        self.use_docint: bool = os.getenv("USE_AZURE_DOCINT", "false").lower() in ("1", "true", "yes")
        self.docint = AzureDocumentIntelligenceService()
        # Diagnostics
        try:
            print(
                f"Azure DocInt configured: use_docint={self.use_docint}, available={getattr(self.docint, 'available', False)}, "
                f"endpoint_set={bool(os.getenv('AZURE_DOCINT_ENDPOINT'))}, key_set={bool(os.getenv('AZURE_DOCINT_KEY'))}"
            )
            if getattr(self.docint, 'error', ""):
                print(f"Azure DocInt init error: {self.docint.error}")
        except Exception:
            pass
    
    def _pdf_has_text(self, pdf_path: str) -> bool:
        """Kiểm tra PDF có text layer không"""
        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    return True
            return False
        except Exception:
            return False
        
    def _ocr_pdf_if_needed(self, pdf_path: str) -> str:
        if self._pdf_has_text(pdf_path):
            print(f"Bỏ qua OCR (đã có text): {pdf_path}")
            return pdf_path

        tmp_path = pdf_path.replace(".pdf", "_ocr.pdf")
        try:
            print(f"[DEBUG] Bắt đầu OCR cho {pdf_path}")
            ocrmypdf.ocr(
                input_file=pdf_path,
                output_file=tmp_path,
                language="vie",
                deskew=True,
                force_ocr=True,
                jobs=4
            )
            print(f"[DEBUG] OCR hoàn tất, file tạm: {tmp_path}")

            # Kiểm tra file tạm có tồn tại và đọc được không
            if not os.path.exists(tmp_path):
                print(f"[ERROR] File OCR {tmp_path} không tồn tại")
                return pdf_path

            # Kiểm tra text sau OCR
            if self._pdf_has_text(tmp_path):
                print(f"[DEBUG] Text layer tìm thấy trong {tmp_path}")
                shutil.move(tmp_path, pdf_path)
                print(f"OCR thành công, cập nhật file: {pdf_path}")
            else:
                print(f"[ERROR] OCR xong nhưng không có text layer: {tmp_path}")
                return tmp_path
        except Exception as e:
            print(f"[ERROR] OCR thất bại cho {pdf_path}: {str(e)}")
            return pdf_path
        return pdf_path

    def _preprocess_text(self, text: str) -> str:
        """Preprocess text for better quality"""
        import re
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text
    
    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            if file_path.lower().endswith(".pdf"):
                file_path = self._ocr_pdf_if_needed(file_path)
                pages = self.extract_text_from_pdf(file_path)
            else:
                pages = [(1, self.extract_text_from_file(file_path))]

            all_chunks = []
            total_chunks = 0

            for page_number, text in pages:
                processed_content = self._preprocess_text(text)
                if not processed_content.strip():
                    continue 

                chunks = self.text_splitter.split_text(processed_content)
                for chunk in chunks:
                    total_chunks += 1
                    embedding = self.azure_openai.create_embedding(chunk)
                    chunk_data = {
                        'title': os.path.splitext(os.path.basename(file_path))[0],
                        'content': chunk,
                        'file_path': f"{file_path}#chunk{total_chunks}",
                        'embedding': embedding,
                        'chunk_index': total_chunks,
                        'total_chunks': None,
                        'doc_metadata': {
                            'chunk_size': len(chunk),
                            'import_timestamp': datetime.datetime.now().isoformat(),
                            'original_file': os.path.basename(file_path),
                            'original_path': file_path,
                            'page_number': page_number
                        }
                    }
                    all_chunks.append(chunk_data)

            # update total_chunks
            for item in all_chunks:
                item['total_chunks'] = total_chunks

            return all_chunks
        except Exception as e:
            print(f"Error processing document: {e}")
            return []


    def extract_text_from_pdf(self, file_path: str) -> List[Tuple[int, str]]:
        pages_text = []
        try:
            if self.use_docint and self.docint.available:
                text = self.docint.extract_text(file_path)
                if text and text.strip():
                    pages_text.append((1, text))
                    print(f"[DEBUG] Azure DocInt extracted text: {len(text)} chars")
                else:
                    print(f"[WARN] Azure DocInt không trích xuất được text từ {file_path}")
                return pages_text

            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                print(f"[DEBUG] Page {page_num}: {len(text)} chars")
                if text.strip():
                    pages_text.append((page_num, text))
                else:
                    print(f"[WARN] Page {page_num} không có text")

            if not pages_text:
                print(f"[ERROR] Không trích xuất được text từ {file_path}")
            return pages_text
        except Exception as e:
            print(f"[ERROR] Lỗi khi trích xuất PDF {file_path}: {str(e)}")
            return []
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file including paragraphs and tables"""
        try:
            if self.use_docint and self.docint.available:
                text = self.docint.extract_text(file_path)
                return text or ""
            doc = DocxDocument(file_path)
            parts = []
            # Paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    parts.append(paragraph.text)
            # Tables
            for table in getattr(doc, 'tables', []):
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = "\n".join(p.text for p in cell.paragraphs if p.text and p.text.strip())
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        parts.append(" \t ".join(row_text))
            text = "\n".join(parts).strip()
            # Optional fallback if empty and docx2txt is available
            if not text:
                try:
                    import docx2txt  # type: ignore
                    text = docx2txt.process(file_path) or ""
                except Exception:
                    pass
            return text
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            # Best-effort fallback using docx2txt if python-docx fails
            try:
                import docx2txt  # type: ignore
                text = docx2txt.process(file_path) or ""
                return text
            except Exception as e2:
                print(f"Fallback docx2txt failed: {e2}")
                return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Document Intelligence can parse tables; when enabled, don't fallback
            if self.use_docint and self.docint.available:
                text = self.docint.extract_text(file_path)
                return text or ""
            df = pd.read_excel(file_path)
            text = ""
            for column in df.columns:
                text += f"{column}: "
                text += " ".join([str(cell) for cell in df[column].dropna()]) + "\n"
            return text
        except Exception as e:
            print(f"Error extracting Excel: {e}")
            return ""
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using Azure Document Intelligence only"""
        try:
            if not (self.use_docint and self.docint.available):
                raise RuntimeError("Azure Document Intelligence is not configured or not available.")
            text = self.docint.extract_text(file_path)
            return text or ""
        except Exception as e:
            print(f"Error extracting image: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self.extract_text_from_image(file_path)
        else:
            return ""
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Compatibility wrapper for existing extract_text method"""
        return self.extract_text(file_path)
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for better processing"""
        if not text.strip():
            return []
        
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.batch_size):
            chunk = " ".join(words[i:i + self.batch_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        
        return chunks
    
    def create_embedding(self, text: str) -> np.ndarray:
        """Create embedding for text using Azure OpenAI"""
        try:
            embedding = self.azure_openai.create_embedding(text)
            return embedding
        except Exception as e:
            print(f"Error creating embedding: {e}")
            return np.zeros(1536)  # Default dimension for text-embedding-3-small
    
    def get_file_hash(self, file_path: str) -> str:
        """Get file hash for version control"""
        try:
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            return file_hash
        except Exception as e:
            print(f"Error getting file hash: {e}")
            return ""
    
    def is_file_updated(self, file_path: str, stored_hash: str) -> bool:
        """Check if file has been updated"""
        current_hash = self.get_file_hash(file_path)
        return current_hash != stored_hash
